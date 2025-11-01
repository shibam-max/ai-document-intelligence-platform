from fastapi import FastAPI, HTTPException, UploadFile, File
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from typing import List, Optional, Dict, Any
import uvicorn
import logging
from datetime import datetime
import asyncio
import os
import io
import fitz  # PyMuPDF
from PIL import Image
import pytesseract
import docx
import pandas as pd
from openpyxl import load_workbook
import json

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

app = FastAPI(
    title="Document Processor Service",
    description="Advanced document processing and text extraction service",
    version="1.0.0"
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Pydantic Models
class ProcessingRequest(BaseModel):
    filename: str
    processing_type: str = "full_extraction"
    extract_images: bool = False
    extract_tables: bool = False
    ocr_enabled: bool = True
    language: str = "eng"

class ProcessingResponse(BaseModel):
    document_id: str
    filename: str
    file_type: str
    text_content: str
    metadata: Dict[str, Any]
    images: List[str] = []
    tables: List[Dict[str, Any]] = []
    processing_time: float
    timestamp: datetime

class DocumentMetadata(BaseModel):
    pages: int
    word_count: int
    character_count: int
    file_size: int
    creation_date: Optional[datetime]
    modification_date: Optional[datetime]
    author: Optional[str]
    title: Optional[str]

# Document Processor
class DocumentProcessor:
    def __init__(self):
        self.supported_formats = {
            '.pdf': self.process_pdf,
            '.docx': self.process_docx,
            '.doc': self.process_doc,
            '.txt': self.process_txt,
            '.xlsx': self.process_excel,
            '.xls': self.process_excel,
            '.csv': self.process_csv,
            '.png': self.process_image,
            '.jpg': self.process_image,
            '.jpeg': self.process_image,
            '.tiff': self.process_image,
            '.bmp': self.process_image
        }
    
    async def process_document(self, file: UploadFile, request: ProcessingRequest) -> ProcessingResponse:
        """Main document processing function"""
        start_time = datetime.utcnow()
        
        # Read file content
        content = await file.read()
        file_extension = os.path.splitext(request.filename)[1].lower()
        
        if file_extension not in self.supported_formats:
            raise HTTPException(status_code=400, detail=f"Unsupported file format: {file_extension}")
        
        try:
            # Process based on file type
            processor = self.supported_formats[file_extension]
            result = await processor(content, request)
            
            # Calculate processing time
            processing_time = (datetime.utcnow() - start_time).total_seconds()
            
            # Generate document ID
            document_id = f"proc_{hash(request.filename + str(start_time))}"
            
            return ProcessingResponse(
                document_id=document_id,
                filename=request.filename,
                file_type=file_extension,
                text_content=result["text"],
                metadata=result["metadata"],
                images=result.get("images", []),
                tables=result.get("tables", []),
                processing_time=processing_time,
                timestamp=datetime.utcnow()
            )
            
        except Exception as e:
            logger.error(f"Document processing failed: {str(e)}")
            raise HTTPException(status_code=500, detail=f"Processing failed: {str(e)}")
    
    async def process_pdf(self, content: bytes, request: ProcessingRequest) -> Dict[str, Any]:
        """Process PDF documents"""
        doc = fitz.open(stream=content, filetype="pdf")
        
        text_content = ""
        images = []
        tables = []
        
        for page_num in range(len(doc)):
            page = doc.load_page(page_num)
            
            # Extract text
            page_text = page.get_text()
            text_content += page_text + "\n"
            
            # Extract images if requested
            if request.extract_images:
                image_list = page.get_images()
                for img_index, img in enumerate(image_list):
                    xref = img[0]
                    pix = fitz.Pixmap(doc, xref)
                    if pix.n - pix.alpha < 4:  # GRAY or RGB
                        img_data = pix.tobytes("png")
                        images.append(f"page_{page_num}_img_{img_index}.png")
                    pix = None
            
            # Extract tables if requested
            if request.extract_tables:
                try:
                    tables_on_page = page.find_tables()
                    for table_index, table in enumerate(tables_on_page):
                        table_data = table.extract()
                        tables.append({
                            "page": page_num,
                            "table_index": table_index,
                            "data": table_data
                        })
                except:
                    pass  # Table extraction might fail
        
        # Extract metadata
        metadata = {
            "pages": len(doc),
            "word_count": len(text_content.split()),
            "character_count": len(text_content),
            "file_size": len(content),
            "title": doc.metadata.get("title"),
            "author": doc.metadata.get("author"),
            "creation_date": doc.metadata.get("creationDate"),
            "modification_date": doc.metadata.get("modDate")
        }
        
        doc.close()
        
        return {
            "text": text_content,
            "metadata": metadata,
            "images": images,
            "tables": tables
        }
    
    async def process_docx(self, content: bytes, request: ProcessingRequest) -> Dict[str, Any]:
        """Process DOCX documents"""
        doc = docx.Document(io.BytesIO(content))
        
        text_content = ""
        tables = []
        
        # Extract paragraphs
        for paragraph in doc.paragraphs:
            text_content += paragraph.text + "\n"
        
        # Extract tables if requested
        if request.extract_tables:
            for table_index, table in enumerate(doc.tables):
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                
                tables.append({
                    "table_index": table_index,
                    "data": table_data
                })
        
        # Extract metadata
        props = doc.core_properties
        metadata = {
            "pages": 1,  # DOCX doesn't have clear page concept
            "word_count": len(text_content.split()),
            "character_count": len(text_content),
            "file_size": len(content),
            "title": props.title,
            "author": props.author,
            "creation_date": props.created,
            "modification_date": props.modified
        }
        
        return {
            "text": text_content,
            "metadata": metadata,
            "tables": tables
        }
    
    async def process_doc(self, content: bytes, request: ProcessingRequest) -> Dict[str, Any]:
        """Process DOC documents (legacy format)"""
        # For DOC files, we'll use a simplified approach
        # In production, you might want to use python-docx2txt or similar
        try:
            import docx2txt
            text_content = docx2txt.process(io.BytesIO(content))
        except:
            text_content = "DOC file processing requires additional libraries"
        
        metadata = {
            "pages": 1,
            "word_count": len(text_content.split()),
            "character_count": len(text_content),
            "file_size": len(content)
        }
        
        return {
            "text": text_content,
            "metadata": metadata
        }
    
    async def process_txt(self, content: bytes, request: ProcessingRequest) -> Dict[str, Any]:
        """Process text files"""
        try:
            text_content = content.decode('utf-8')
        except UnicodeDecodeError:
            try:
                text_content = content.decode('latin-1')
            except:
                text_content = content.decode('utf-8', errors='ignore')
        
        metadata = {
            "pages": 1,
            "word_count": len(text_content.split()),
            "character_count": len(text_content),
            "file_size": len(content)
        }
        
        return {
            "text": text_content,
            "metadata": metadata
        }
    
    async def process_excel(self, content: bytes, request: ProcessingRequest) -> Dict[str, Any]:
        """Process Excel files"""
        try:
            # Try with openpyxl first (for .xlsx)
            workbook = load_workbook(io.BytesIO(content))
            text_content = ""
            tables = []
            
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                sheet_data = []
                
                for row in sheet.iter_rows(values_only=True):
                    row_data = [str(cell) if cell is not None else "" for cell in row]
                    sheet_data.append(row_data)
                    text_content += " ".join(row_data) + "\n"
                
                if request.extract_tables:
                    tables.append({
                        "sheet_name": sheet_name,
                        "data": sheet_data
                    })
            
        except:
            # Fallback to pandas for .xls files
            try:
                df = pd.read_excel(io.BytesIO(content))
                text_content = df.to_string()
                
                if request.extract_tables:
                    tables = [{
                        "sheet_name": "Sheet1",
                        "data": df.values.tolist()
                    }]
            except Exception as e:
                raise HTTPException(status_code=500, detail=f"Excel processing failed: {str(e)}")
        
        metadata = {
            "pages": len(tables) if tables else 1,
            "word_count": len(text_content.split()),
            "character_count": len(text_content),
            "file_size": len(content)
        }
        
        return {
            "text": text_content,
            "metadata": metadata,
            "tables": tables
        }
    
    async def process_csv(self, content: bytes, request: ProcessingRequest) -> Dict[str, Any]:
        """Process CSV files"""
        try:
            text_content = content.decode('utf-8')
            df = pd.read_csv(io.StringIO(text_content))
            
            # Convert to text
            text_content = df.to_string()
            
            tables = []
            if request.extract_tables:
                tables = [{
                    "sheet_name": "CSV_Data",
                    "data": df.values.tolist()
                }]
            
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"CSV processing failed: {str(e)}")
        
        metadata = {
            "pages": 1,
            "word_count": len(text_content.split()),
            "character_count": len(text_content),
            "file_size": len(content),
            "rows": len(df),
            "columns": len(df.columns)
        }
        
        return {
            "text": text_content,
            "metadata": metadata,
            "tables": tables
        }
    
    async def process_image(self, content: bytes, request: ProcessingRequest) -> Dict[str, Any]:
        """Process image files with OCR"""
        if not request.ocr_enabled:
            return {
                "text": "OCR disabled for image processing",
                "metadata": {"file_size": len(content)}
            }
        
        try:
            # Open image
            image = Image.open(io.BytesIO(content))
            
            # Perform OCR
            text_content = pytesseract.image_to_string(image, lang=request.language)
            
            metadata = {
                "pages": 1,
                "word_count": len(text_content.split()),
                "character_count": len(text_content),
                "file_size": len(content),
                "image_width": image.width,
                "image_height": image.height,
                "image_mode": image.mode
            }
            
            return {
                "text": text_content,
                "metadata": metadata
            }
            
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Image OCR failed: {str(e)}")

# Initialize processor
processor = DocumentProcessor()

@app.on_event("startup")
async def startup_event():
    """Initialize service on startup"""
    logger.info("Starting Document Processor Service...")
    logger.info("Document Processor Service started successfully")

@app.get("/health")
async def health_check():
    """Health check endpoint"""
    return {
        "status": "healthy",
        "timestamp": datetime.utcnow(),
        "supported_formats": list(processor.supported_formats.keys())
    }

@app.post("/process", response_model=ProcessingResponse)
async def process_document(
    file: UploadFile = File(...),
    processing_type: str = "full_extraction",
    extract_images: bool = False,
    extract_tables: bool = False,
    ocr_enabled: bool = True,
    language: str = "eng"
):
    """Process uploaded document"""
    request = ProcessingRequest(
        filename=file.filename,
        processing_type=processing_type,
        extract_images=extract_images,
        extract_tables=extract_tables,
        ocr_enabled=ocr_enabled,
        language=language
    )
    
    logger.info(f"Processing document: {file.filename}")
    result = await processor.process_document(file, request)
    logger.info(f"Document processed successfully: {result.document_id}")
    
    return result

@app.get("/formats")
async def supported_formats():
    """Get supported file formats"""
    return {
        "supported_formats": list(processor.supported_formats.keys()),
        "ocr_languages": ["eng", "fra", "deu", "spa", "ita", "por"]
    }

@app.get("/metrics")
async def get_metrics():
    """Get service metrics"""
    return {
        "supported_formats": len(processor.supported_formats),
        "service_status": "running"
    }

if __name__ == "__main__":
    uvicorn.run(
        "main:app",
        host="0.0.0.0",
        port=8000,
        reload=True,
        log_level="info"
    )